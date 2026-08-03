#!/usr/bin/env python3
"""Generate the system-focused AEGIS SOC Dashboard project book."""
from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Keep the downloadable filename simple; parentheses in the old upload made the
# attachment awkward to retrieve from some GitHub/chat clients.
OUT = Path("AEGIS_SOC_Dashboard_Project_Book.docx")
doc = Document()
section = doc.sections[0]
section.page_width, section.page_height = Cm(21), Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin = section.bottom_margin = Cm(2.3)

for name, size, bold in [
    ("Normal", 12, False), ("Title", 14, True),
    ("Heading 1", 14, True), ("Heading 2", 14, True),
]:
    style = doc.styles[name]
    style.font.name = "Aptos" if name != "Normal" else "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)


def page_break():
    doc.add_page_break()


def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(7)
    for r in p.runs:
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def prose(text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.35
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.8)
    p.add_run(text)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    return p


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.add_run(text)
    for r in p.runs:
        r.font.name = "Courier New"
        r.font.size = Pt(12)
    ppr = p._element.get_or_add_pPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), "F2F4F7")
    ppr.append(shade)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "7")
        border.set(qn("w:color"), "7A8793")
        border.set(qn("w:space"), "5")
        borders.append(border)
    ppr.append(borders)
    return p


def image(path, label):
    if not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(15.5))
    c = doc.add_paragraph(label)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in c.runs:
        r.italic = True
        r.font.size = Pt(12)


# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(90)
r = p.add_run("AEGIS SOC DASHBOARD")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0, 0, 0)
p = doc.add_paragraph("System Book")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(14); p.runs[0].bold = True
p = doc.add_paragraph("Real-Time Monitoring, Threat Detection, and Coordinated Defense\nfor the GNS3 AEGIS-SecureCompany Lab")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in p.runs: r.font.size = Pt(12)
p = doc.add_paragraph("Current implemented system only")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
p.runs[0].italic = True
p.runs[0].font.size = Pt(12)
page_break()

# TOC — only table in book
topics = [
    "System at a Glance", "What the System Monitors", "Current Network Topology",
    "System Components and Responsibilities", "How a Security Event Moves Through AEGIS",
    "Detection Sources", "Dashboard Functions", "Defense Process",
    "AI Analysis, Reports, and Telegram Alerts", "Database and API Design",
    "Authentication and Security Controls", "Deployment and Configuration",
    "How the System Was Built and Verified", "Test Scenarios and Expected Results",
    "Current Boundaries and Limitations", "Essential Terms", "Command Reference",
]
heading("TABLE OF CONTENTS")
t = doc.add_table(rows=1, cols=1)
t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = t.cell(0, 0); cell.text = "System Topics"
cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
shade = OxmlElement("w:shd"); shade.set(qn("w:fill"), "E7E7E7"); cell._tc.get_or_add_tcPr().append(shade)
for r in cell.paragraphs[0].runs:
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 0, 0)
for topic in topics:
    row = t.add_row(); row.cells[0].text = topic
    for r in row.cells[0].paragraphs[0].runs: r.font.name = "Times New Roman"; r.font.size = Pt(12)
page_break()

heading("SYSTEM AT A GLANCE")
prose("AEGIS is a Security Operations Center dashboard for the GNS3 AEGIS-SecureCompany lab. It collects real security events from the lab, stores and streams them through a cloud API, shows them to an administrator, and records defense actions. The dashboard is not a simulation engine: attacks and host-level defenses happen on the GNS3 virtual machines; the web application provides central visibility, analysis, alerts, and administrator control.")
prose("An attack begins from Kali Linux and travels through MikroTik and pfSense toward a protected company server. Suricata, Fail2ban, SSH logs, and HTTP security logs produce evidence. The AEGIS forwarder reads that evidence, normalizes it, and sends it to the Express API using an ingest key. The API stores it in PostgreSQL, broadcasts it to the React dashboard with Server-Sent Events, and can notify Telegram or evaluate a defense response.")
prose("The implemented v4 lab protects four company workloads: an Apache web server and BIND9 DNS server in the DMZ, plus a MySQL customer database and OpenLDAP server in the internal zone. Two OVS switches connect the service pairs to pfSense. A dedicated management-zone VM runs the hub forwarder. Mail, VoIP, CCTV, Cowrie, the removed second MikroTik router, and the old shared-switch topology are not part of the current system and are intentionally excluded.")

heading("WHAT THE SYSTEM MONITORS")
prose("AEGIS answers four practical questions for the administrator: What happened? Where did it happen? How serious is it? What response was taken? It combines events that would otherwise remain separated across firewall and server logs.")
bullet("The company web server at 10.10.10.10, including Apache/HTTP activity, SSH authentication, Suricata alerts, and Fail2ban actions.")
bullet("The company customer database at 10.20.20.10, including MySQL activity, SSH authentication, Suricata alerts, and Fail2ban actions.")
bullet("The BIND9 DNS server at 10.10.10.20 and OpenLDAP server at 10.20.20.20, including their service, SSH, and Fail2ban evidence.")
bullet("pfSense at 10.0.23.2 on the WAN link and 10.10.10.1, 10.20.20.1, and 10.30.30.1 on the protected zones.")
bullet("The AEGIS forwarder VM at 10.30.30.10, which collects remote logs and reports component health.")

heading("CURRENT NETWORK TOPOLOGY")
prose("MikroTik CHR connects the external/NAT network, the Kali attacker network, and the point-to-point WAN link to pfSense. Kali receives a dynamic 192.168.10.x address from MikroTik. pfSense separates the DMZ, internal-services, and management networks. Public-Services and Internal-Services OVS switches attach the two servers in each protected zone. This separation allows attacks to be observed at the boundary while keeping each protected service in the correct security zone.")
code("Internet/NAT\n    |\nMikroTik CHR\n    |-- Kali attacker: DHCP 192.168.10.x\n    `-- WAN link 10.0.23.1/30\n             |\n        pfSense 10.0.23.2/30\n             |-- DMZ  10.10.10.0/24 -> OVS -> web .10, DNS .20\n             |-- INT  10.20.20.0/24 -> OVS -> MySQL .10, LDAP .20\n             `-- MGMT 10.30.30.0/24 -> AEGIS forwarder 10.30.30.10")
prose("Kali needs a route through MikroTik before it can reach the protected 10.0.0.0/8 networks:")
code("sudo ip route add 10.0.0.0/8 via 192.168.10.1")

heading("SYSTEM COMPONENTS AND RESPONSIBILITIES")
prose("Kali Linux generates controlled attacks for validation. MikroTik provides routing and DHCP on the attacker side. pfSense is the network boundary and firewall between the attacker and the protected zones.", "Lab network. ")
prose("The web, DNS, MySQL customer-database, and LDAP servers run the protected services and security sensors. Their logs are the original source of truth for activity on each host.", "Protected workloads. ")
prose("The Python hub agent runs on the AEGIS VM. It connects to the protected machines over SSH, tails supported logs, converts different log formats into one event structure, sends events to the API, and polls for work needed by the defense workflow.", "Forwarder. ")
prose("The Express 5 and TypeScript API validates requests, authenticates ingest traffic, reads and writes data through Drizzle ORM, broadcasts live events, schedules reports, integrates AI and Telegram, and exposes the routes used by the dashboard.", "API server. ")
prose("Supabase PostgreSQL stores security events, alerts, system status, reports, network hosts, blocked IPs, defense actions, and defense rules. The React/Vite dashboard reads this data and receives live changes over SSE.", "Data and presentation. ")

heading("HOW A SECURITY EVENT MOVES THROUGH AEGIS")
prose("First, a sensor or operating-system log records an event on a protected VM. Second, the hub forwarder reads the new line over SSH and converts source-specific fields into a consistent event containing time, source, event type, severity, source IP, destination IP, message, and raw evidence. Third, it sends an authenticated POST request to an ingest endpoint.")
prose("The API validates the body, saves the event, updates related alert state where applicable, and publishes the event to connected dashboard clients. The dashboard updates without a refresh because the browser maintains an SSE connection to the API. High-severity activity can also trigger Telegram notification and defense evaluation.")
code("Sensor/log -> Forwarder -> X-AEGIS-Key POST -> Express API\n           -> Drizzle ORM -> Supabase PostgreSQL\n           -> SSE -> React dashboard\n           -> alert / Telegram / defense decision")

heading("DETECTION SOURCES")
prose("Suricata examines network traffic and produces structured EVE JSON alerts for signatures and suspicious behavior. AEGIS preserves the useful network evidence, such as source, destination, protocol, signature, and severity.", "Suricata. ")
prose("Fail2ban watches repeated failures and applies host firewall bans. AEGIS records ban and unban actions so the administrator can see that a host-level automatic response occurred.", "Fail2ban. ")
prose("Authentication monitoring records successful and failed SSH access. Repeated failures are important for brute-force detection, while successful access is retained for audit context.", "SSH. ")
prose("HTTP monitoring records web requests and web-attack evidence. This supports visibility for scanning, injection attempts, and abnormal request volume against the Apache service.", "HTTP. ")
prose("These sources have different formats, so normalization in the forwarder is essential. The dashboard does not need separate rendering logic for every original log format.")

heading("DASHBOARD FUNCTIONS")
prose("The Command Center summarizes attack volume, severity, active threats, recent events, and component health. Security Events provides the detailed live feed. Active Alerts supports acknowledge and resolve actions. Connections presents SSH and HTTP connection evidence.")
prose("Network Monitor shows the current pfSense, web, DNS, database, LDAP, and forwarder topology with observed hosts and last-seen information. Defense Center shows native automatic protection, blocked addresses, defense history, and administrator block or unblock controls.")
prose("Reports presents scheduled and generated SOC summaries. AI tools explain individual events and provide defense recommendations. Defense Rules configures event-response matching, Attack Flow explains the live pipeline, and Settings controls supported application options.")
image("figure_4_9.png", "AEGIS dashboard overview")

heading("DEFENSE PROCESS")
prose("AEGIS uses a hybrid defense model. Fast host-level automatic defense remains close to the protected machine: for example, Fail2ban detects repeated SSH failures and adds an iptables block. The forwarder reports the result so that the dashboard and audit history match the real VM state.")
prose("For manual defense, an authenticated administrator enters an IP address and reason in Defense Center. The API validates the request, prevents unsafe or whitelisted targets from being used, records the action, and coordinates the requested block. For a pfSense target, the hub forwarder uses its configured SSH key to run the restricted pfSense firewall command; the legacy REST settings are not required. Unblock follows the same controlled and audited path.")
prose("A defense record is not proof by itself that a firewall command succeeded. The system keeps command/result or action status so the interface can distinguish a request from an executed response. This is important because the web dashboard must not pretend to perform a network action that failed in the lab.")

heading("AI ANALYSIS, REPORTS, AND TELEGRAM ALERTS")
prose("Groq AI is an optional explanation layer. It receives structured security context from the API and turns it into a human-readable event explanation, threat briefing, report summary, or IP-defense recommendation. It does not replace Suricata or Fail2ban and it is not the source of the raw event. When GROQ_API_KEY is absent, core event collection and monitoring continue without AI output.")
prose("The reporting scheduler summarizes stored activity for a selected period. Telegram is an optional delivery channel for important alerts and reports. TELEGRAM_BOT_TOKEN and TELEGRAM_CHAT_ID must be configured; otherwise dashboard monitoring continues without Telegram delivery.")
prose("AI recommendations must be treated as analyst support. Firewall decisions still require validation, whitelisting, command sanitization, and auditable execution.")

heading("DATABASE AND API DESIGN")
prose("The system uses PostgreSQL through Drizzle ORM. TypeScript schema definitions are the source used by application queries. The important data groups are security events, alerts, component status and network hosts, reports, and defense rules/actions/blocked IPs.")
prose("Sensor traffic uses dedicated ingest routes protected by the X-AEGIS-Key header. Browser-facing routes provide events, alerts, system status, network state, defense, reports, and AI functions. The SSE route is one-way from server to browser, which fits monitoring because the dashboard needs immediate updates but does not need a permanent two-way socket.")
code("POST /api/ingest/event\nPOST /api/ingest/suricata\nPOST /api/ingest/fail2ban\nPOST /api/ingest/ssh\nPOST /api/ingest/http\nPOST /api/network/hosts\nGET  /api/events/stream")

heading("AUTHENTICATION AND SECURITY CONTROLS")
prose("Ingest authentication and administrator authentication are separate. Sensors use AEGIS_INGEST_KEY, while privileged dashboard operations use session/JWT and administrator controls. Secrets belong in environment variables and must not be embedded in source code or screenshots.")
bullet("Validate every ingest body before saving it.")
bullet("Use a different strong value for the ingest key, admin key, and session secret.")
bullet("Sanitize defense targets and never allow loopback, internal management, or whitelisted addresses to be blocked accidentally.")
bullet("Keep an audit record for block, unblock, and pfSense actions.")
bullet("Expose only the required API routes and use HTTPS for hosted traffic.")

heading("DEPLOYMENT AND CONFIGURATION")
prose("The repository is a pnpm workspace. The React dashboard is deployed to Vercel, the Express API to Render, and PostgreSQL is hosted by Supabase. The forwarder remains inside the GNS3 management network because it needs SSH reachability to the protected VMs.")
prose("The minimum API configuration is:")
code("SUPABASE_DB_URL=postgresql://user:password@host:6543/postgres\nSESSION_SECRET=<strong-random-secret>\nAEGIS_INGEST_KEY=<sensor-ingest-secret>\nAEGIS_ADMIN_KEY=<administrator-secret>\n\n# Optional integrations\nGROQ_API_KEY=<groq-key>\nTELEGRAM_BOT_TOKEN=<telegram-token>\nTELEGRAM_CHAT_ID=<chat-id>")
prose("Install and start the development services with:")
code("pnpm install\nPORT=3000 pnpm --filter @workspace/api-server run dev\npnpm --filter @workspace/aegis-dashboard run dev")

heading("HOW THE SYSTEM WAS BUILT AND VERIFIED")
prose("Research methodology means the practical method used to understand, build, and verify the system. For AEGIS, the work started by defining the current attack path and protected assets, then identifying the evidence each sensor produces. The API and database contract were designed around that evidence. The forwarder, dashboard, alerts, and defense workflow were integrated afterward. Finally, controlled Kali attacks were used to compare the original logs, stored events, dashboard output, notifications, and firewall result.")
prose("Literature review is an academic name for studying concepts and existing tools before making design choices. In this project it means understanding SOC monitoring, Suricata, Fail2ban, SSE, firewall control, and event normalization. Only concepts implemented by AEGIS matter here; unrelated security technologies are excluded.")
prose("Related work means looking at systems that solve a similar problem and learning from their approach. It does not mean copying their code. For this system, the useful comparison is limited to centralized event monitoring and response; this book focuses on the implemented AEGIS flow rather than describing other products.")

heading("TEST SCENARIOS AND EXPECTED RESULTS")
prose("An SSH brute-force test should produce authentication failures, a Fail2ban ban when its threshold is reached, an AEGIS security event, a live dashboard update, and an auditable defense action. A port scan should be visible in Suricata evidence and the event feed. A web-attack test should appear through HTTP or Suricata monitoring with the attacker and target identified.")
prose("A test passes only when the evidence agrees across layers: the source VM log exists, the API accepted and stored the normalized event, SSE delivered it to the dashboard, and any claimed block can be confirmed on the responsible firewall. Telegram and AI are tested separately because they are optional external integrations.")
prose("All attacks must remain inside the authorized GNS3 lab. Do not run Hydra, sqlmap, hping3, Metasploit, or scanning commands against public or third-party systems.")

heading("CURRENT BOUNDARIES AND LIMITATIONS")
prose("The current topology contains an Apache web server, BIND9 DNS server, MySQL customer database, OpenLDAP server, two zone OVS switches, one forwarder, pfSense, MikroTik, and Kali. Earlier mail, VoIP, CCTV, Cowrie, second-router, and old shared-switch designs are not current implementation components.")
prose("The dashboard depends on network connectivity between the forwarder and protected VMs and between the forwarder and hosted API. Render sleep/restart behavior, external AI or Telegram availability, incorrect SSH permissions, rotated secrets, and incorrect pfSense SSH configuration can interrupt optional or remote actions.")
prose("AEGIS is a focused lab SOC and not a replacement for an enterprise SIEM or a staffed incident-response program. Its results demonstrate an end-to-end monitoring and defense workflow in the defined environment.")

heading("ESSENTIAL TERMS")
prose("Security Operations Center: the central place used to monitor events, investigate threats, and coordinate response.", "SOC — ")
prose("Intrusion Detection System: a sensor that identifies suspicious traffic or behavior and raises an alert.", "IDS — ")
prose("Security Information and Event Management: centralized collection, storage, search, and correlation of security events. AEGIS implements a focused subset for its lab.", "SIEM — ")
prose("Server-Sent Events: a long-lived HTTP connection used by the API to push new events to the browser.", "SSE — ")
prose("the process of converting Suricata, Fail2ban, SSH, and HTTP evidence into a consistent event shape.", "Normalization — ")
prose("a network segment exposed to controlled external access while separated from the internal database and management zones.", "DMZ — ")

heading("COMMAND REFERENCE")
prose("Run repository checks:")
code("pnpm run typecheck\npnpm run build")
prose("Apply the Drizzle schema to the configured database:")
code("pnpm --filter @workspace/db run push")
prose("Regenerate API clients after changing the OpenAPI contract:")
code("pnpm --filter @workspace/api-spec run codegen")
prose("Start the hub forwarder only after its API URL, ingest key, SSH targets, and key permissions are configured. Run the current hub mode with the local configuration file beside the script:")
code("cd scripts/src\npython3 aegis_forwarder.py --mode hub")

page_break()
heading("SYSTEM SUMMARY")
prose("AEGIS joins a real GNS3 attack path, host and network sensors, a Python hub forwarder, an authenticated TypeScript API, PostgreSQL storage, live SSE delivery, a React dashboard, optional AI and Telegram services, and audited defense controls. Reading the book from the current topology through the event and defense flows explains where evidence originates, how it reaches the dashboard, and which component is responsible for each response.")
prose("The generated document intentionally contains no unrelated academic filler and no obsolete mail, VoIP, CCTV, Cowrie, second-router, or old shared-switch implementation sections.")

doc.save(OUT)
print(f"Saved {OUT} ({OUT.stat().st_size / 1024:.1f} KiB)")
